import streamlit as st
import os
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
import re
import hashlib
from typing import TypedDict, Optional, List, Dict
from langchain_google_genai import ChatGoogleGenerativeAI
from langgraph.graph import StateGraph, END
from rag_pipeline import run_rag_pipeline
from database import init_db, login_user, register_user, save_quiz, get_quiz_history
from docx import Document as DocxDocument

# ------------------- Session State Initialization -------------------
if "user" not in st.session_state:
    st.session_state.user = None
if "quiz_data" not in st.session_state:
    st.session_state.quiz_data = []
if "use_text" not in st.session_state:
    st.session_state.use_text = False
if "current_step" not in st.session_state:
    st.session_state.current_step = "upload_node"
if "state" not in st.session_state:
    st.session_state.state = {
        "file": None,
        "raw_text": "",
        "context_text": "",
        "manual_topic": "",
        "file_hash": "",
        "quiz_data": [],
        "num_mcqs": 5
    }

# ------------------- LLM Setup -------------------
google_api_key = os.environ.get("GOOGLE_API_KEY")
try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key=google_api_key
    )
except Exception:
    st.error("⚠️ GOOGLE_API_KEY not found. LLM features disabled.")
    llm = None

# ------------------- Streamlit Page & Custom Styling -------------------
st.set_page_config(page_title="🧠 AI Quiz Generator", layout="wide")
init_db()

# =========================
# 🔥 UPDATED SIDEBAR + UI
# =========================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;600;700;800;900&display=swap');

:root {
    --primary-accent: #33CCFF;
    --secondary-dark: #0B1A2B;
    --background-main: #162231;
    --card-bg-light: #162231;
    --text-color: #F0F5FA;
    --muted-text: #9FB8C9;
}

body, .stApp {
    font-family: 'Space Grotesk', sans-serif;
    color: var(--text-color);
    background: var(--background-main);
    overflow-x: hidden;
}

header { visibility: hidden; height: 0; }

/* ============================= */
/* SIDEBAR STYLE + RESPONSIVE */
/* ============================= */
section[data-testid="stSidebar"] {
    width: 380px !important;
    min-width: 380px !important;
    max-width: 380px !important;
    resize: none !important;
    background: linear-gradient(180deg, #0B1A2B, #162231) !important;
    border-right: 2px solid #33CCFF;
    box-shadow: 0 0 25px rgba(0,255,224,0.3), 0 0 60px rgba(51,204,255,0.15) inset;
    backdrop-filter: blur(14px);
}

section[data-testid="stSidebar"] > div:first-child {
    resize: none !important;
    overflow: auto !important;
}

@media (max-width: 800px) {
    section[data-testid="stSidebar"] {
        width: 300px !important;
        min-width: 300px !important;
        max-width: 300px !important;
    }
}

/* ============================= */
/* Sidebar header animated neon */
/* ============================= */
.sidebar-header-animated {
    font-size: 1.5em;
    font-weight: 600;
    color: #E9F6FB;
    text-align: center;
    padding: 8px 0;
    margin-bottom: 12px;
    text-shadow: 0 0 4px #66e0ff, 0 0 4px #99f0ff;
    animation: neonPulse 1.5s ease-in-out infinite alternate;
}

@keyframes neonPulse {
    0% { text-shadow: 0 0 2px #66e0ff, 0 0 6px #99f0ff; }
    100% { text-shadow: 0 0 10px #00FFE0, 0 0 20px #33CCFF; }
}

/* ============================= */
/* Main H1 & P Responsive */
/* ============================= */
h1 {
    text-align: center;
    font-size: 3.0em;
    font-weight: 900;
    background: linear-gradient(90deg, #33CCFF, #00FFE0);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 8px;
    text-shadow: 0 0 10px rgba(51, 204, 255, 0.9), 0 0 25px rgba(0, 255, 224, 0.7);
}

p {
    text-align: center;
    color: var(--muted-text);
    font-size: 1em;
    margin-bottom: 15px;
}

/* ============================= */
/* QUIZ CARD */
/* ============================= */
.quiz-card {
    background-color: var(--card-bg-light);
    padding: 20px 20px;
    border-radius: 16px;
    margin: 15px auto;
    border-left: 6px solid var(--primary-accent);
    box-shadow: 0 0 15px rgba(0, 0, 0, 0.6), 0 0 25px rgba(51, 204, 255, 0.2);
}

.quiz-question {
    font-size: 1.2em;
    font-weight: 600;
    margin-bottom: 12px;
    color: var(--primary-accent);
}

.quiz-option {
    padding: 10px 15px;
    background-color: var(--secondary-dark);
    border-radius: 10px;
    margin-bottom: 8px;
    cursor: pointer;
    color: var(--text-color);
    font-size: 0.95em;
}

/* ============================= */
/* BUTTON STYLE */
/* ============================= */
.stDownloadButton button, 
.stButton button { 
    background: linear-gradient(135deg, #0B1A2B, #162231) !important;
    color: #33CCFF !important;
    font-weight: 700;
    border-radius: 18px !important;
    border: 2px solid #33CCFF !important;
}

/* ============================= */
/* LOGIN SCREEN - RESPONSIVE */
/* ============================= */
.right-login-box {
    display: flex;
    justify-content: center;
    align-items: center;
    position: fixed; 
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    background-color: rgba(22, 34, 49, 0.0); 
    z-index: 9999; 
    pointer-events: none;
    transform: translateX(200px);
}

.animated-message {
    background: linear-gradient(135deg, #162231, #0B1A2B);
    color: #33CCFF;
    border: 2px solid #00FFE0;
    padding: 70px 100px;
    border-radius: 22px;
    font-size: 2.5em;
    font-weight: 800;
    max-width: 90%;
    text-align: center;
    box-shadow: 0 0 40px rgba(0, 255, 224, 0.5), 0 0 20px rgba(51, 204, 255, 0.8);
    animation: fadein 1s, pulse 2s infinite alternate;
}

@media (max-width: 800px) {
    .right-login-box { transform: translateX(0px); }
    .animated-message {
        padding: 40px 20px; 
        font-size: 1.8em;
    }

    .quiz-card { padding: 15px 10px; }
    .quiz-question { font-size: 1em; }
    .quiz-option { font-size: 0.85em; padding: 8px 10px; }
    h1 { font-size: 2.5em; }
    p { font-size: 0.9em; }
}

/* ============================= */
/* Floating Bubbles Mobile */
@media (max-width: 800px) {
    .bubble { width: 15px !important; height: 15px !important; opacity:0.1; }
}

</style>


<div id="bubbles"></div>

<script>
const colors = ['#33CCFF','#00FFE0','#FF33CC','#FFDD33','#33FFAA'];
for(let i=0;i<25;i++){
    const bubble = document.createElement('div');
    bubble.className = 'bubble';
    let size = Math.random()*45+10;
    bubble.style.width = size+'px';
    bubble.style.height = size+'px';
    bubble.style.left = Math.random()*100+'vw';
    bubble.style.background = colors[Math.floor(Math.random()*colors.length)];
    bubble.style.animationDuration = (Math.random()*15+15)+'s';
    bubble.style.animationDelay = (Math.random()*5)+'s';
    document.getElementById('bubbles').appendChild(bubble);
}
</script>
""", unsafe_allow_html=True)


# ------------------- Sidebar: Authentication -------------------
st.sidebar.markdown(
    '<div class="sidebar-header-animated">Authentication</div>',
    unsafe_allow_html=True
)
if st.session_state.user:
    st.sidebar.success(f"Hello {st.session_state.user['username']}")
    if st.sidebar.button("🚪 Logout"):
        st.session_state.user = None
        st.session_state.quiz_data = []
        st.session_state.current_step = "upload_node"
        st.rerun()
else:
    auth_mode = st.sidebar.radio("Choose action", ["Login", "Register"])
    if auth_mode == "Login":
        username = st.sidebar.text_input("Username")
        password = st.sidebar.text_input("Password", type="password")
        if st.sidebar.button("Login"):
            user = login_user(username, password)
            if user:
                st.session_state.user = user
                # Fetch history upon successful login
                st.session_state.quiz_data = get_quiz_history(user["id"]) or []
                st.rerun()
            else:
                st.sidebar.error("Invalid username or password.")
    else:  # Register
        new_username = st.sidebar.text_input("Create username")
        new_password = st.sidebar.text_input("Create password", type="password")
        if st.sidebar.button("Register"):
            success = register_user(new_username, new_password)
            if success:
                st.sidebar.success("Account created! You can now log in.")
            else:
                st.sidebar.error("Username already exists.")

# ------------------- Sidebar: Quiz History -------------------
if st.session_state.user:
    with st.sidebar.expander("📜 Quiz History", expanded=False):
        history = get_quiz_history(st.session_state.user["id"])
        if not history:
            st.info("No quizzes yet.")
        else:
            # Display history in a compact, readable format
            for idx, entry in enumerate(history, 1):
                st.markdown(f"**Quiz {idx}** - {entry['created_at']}")
                for q_num, q in enumerate(entry["quiz_data"], 1):
                    # Use a smaller font for options in the sidebar history
                    st.markdown(f"Q{q_num}. {q['question']}")
                    st.markdown(f"<p style='font-size:0.8em; margin-bottom: 0;'>A) {q['options']['A']}</p>",
                                unsafe_allow_html=True)
                    st.markdown(f"<p style='font-size:0.8em; margin-bottom: 0;'>B) {q['options']['B']}</p>",
                                unsafe_allow_html=True)
                    st.markdown(f"<p style='font-size:0.8em; margin-bottom: 0;'>C) {q['options']['C']}</p>",
                                unsafe_allow_html=True)
                    st.markdown(f"<p style='font-size:0.8em; margin-bottom: 0;'>D) {q['options']['D']}</p>",
                                unsafe_allow_html=True)
                    st.markdown(f"**Answer:** {q['answer']}")
                st.markdown("---")

# ------------------- Main App -------------------
if st.session_state.user:
    # Use the styled H1 and P from the first script
    st.markdown("<h1>🧠 AI Quiz Generator</h1>", unsafe_allow_html=True)
    st.markdown("<p>Upload a PDF/DOCX or type a topic to generate MCQs (Documents are saved per user)</p>",
                unsafe_allow_html=True)


    class QuizState(TypedDict, total=False):
        file: Optional[any]
        raw_text: str
        context_text: str
        manual_topic: str
        file_hash: str
        quiz_data: List[Dict]
        num_mcqs: int


    graph = StateGraph(QuizState)


    # ------------------- Upload Node -------------------
    def upload_node(state: QuizState) -> QuizState:
        def toggle_use_text():
            st.session_state.use_text = not st.session_state.use_text

        # Keep the checkbox in the correct location
        use_text = st.checkbox("Or type a topic manually", value=st.session_state.use_text, on_change=toggle_use_text)
        file = None
        text_input = ""

        if not st.session_state.use_text:
            file = st.file_uploader("📂 Upload PDF or DOCX", type=["pdf", "docx"])
        else:
            text_input = st.text_area("Enter topic or text", height=150, key="manual_text_input")

        num_mcqs = st.number_input("Number of MCQs", 1, 50, state.get("num_mcqs", 5), 1)
        return {"file": file, "manual_topic": text_input, "num_mcqs": int(num_mcqs)}


    # ------------------- Extract Text Node -------------------
    def extract_text_node(state: QuizState) -> QuizState:
        file = state.get("file")
        manual_topic = state.get("manual_topic", "")
        raw_text = ""
        file_hash = ""

        if file:
            file_bytes = file.read()
            file_hash = hashlib.sha256(file_bytes).hexdigest()
            file.seek(0)
            if file.name.lower().endswith(".pdf"):
                reader = PdfReader(file)
                raw_text = "\n".join([p.extract_text() or "" for p in reader.pages])
            elif file.name.lower().endswith(".docx"):
                doc = Document(BytesIO(file_bytes))
                raw_text = "\n".join([p.text for p in doc.paragraphs])
        elif manual_topic:
            raw_text = manual_topic
            file_hash = hashlib.sha256(manual_topic.encode('utf-8')).hexdigest()

        if not raw_text:
            # Use st.warning instead of st.error here, as the graph execution might stop
            st.warning("Please upload a file or enter a topic.")

        return {
            "file": file,
            "raw_text": raw_text,
            "manual_topic": manual_topic,
            "file_hash": file_hash,
            "num_mcqs": state.get("num_mcqs", 5)
        }


    # ------------------- Generate Quiz Node -------------------
    def generate_quiz_node(state: QuizState) -> QuizState:

        raw_text = state.get("raw_text", "").strip()
        manual_topic = state.get("manual_topic", "").strip()
        file_hash = state.get("file_hash", "manual_topic_no_hash")
        num_mcqs = state.get("num_mcqs", 5)

        if not raw_text or not llm:
            # Use st.warning instead of st.error here, as the graph execution might stop
            st.warning("LLM not initialized or empty input. Skipping quiz generation.")
            return {"raw_text": raw_text, "context_text": "", "num_mcqs": num_mcqs, "quiz_data": []}

        query = manual_topic if manual_topic else raw_text[:100]

        # ---------------- RAG Pipeline ----------------
        context_text = run_rag_pipeline(raw_text, query, file_hash)
        if not context_text:
            st.warning("RAG pipeline returned empty context. Using first 1000 chars as fallback.")
            context_text = raw_text[:1000]

        # ---------------- LLM Quiz Generation ----------------
        prompt = f"""
            Generate {num_mcqs} multiple-choice questions (MCQs) from the following text.
            generate mcqs in that language which user tells you by default generate quiz in English language
            Format each question clearly with options and mark the correct answer at the end.

            CRITICAL INSTRUCTIONS:
            1. Do NOT use phrases like "provided text," "given text," "as in the text," or "According to the text."
            2. Instead, refer to the actual topic or subject matter.
            3. Start directly with the first question, no extra introductory text.

            Example format:
            Q1. What is AI?
            A) Option 1
            B) Option 2
            C) Option 3
            D) Option 4
            Answer: B

        Text:
        {context_text}
        """

        try:
            result = llm.invoke(prompt)
            output_text = getattr(result, "content", None) or getattr(result, "output_text", "")
        except Exception as e:
            st.error(f"LLM Error: {e}")
            return {"raw_text": raw_text, "context_text": context_text, "num_mcqs": num_mcqs, "quiz_data": []}

        output_text = re.sub(r'(?i)question\s*\d*[:.]', lambda m: f"Q", output_text)
        output_text = output_text.replace("Option ", "").replace("Answer:", "Answer:")

        pattern = r"""Q\d*[\.\)]?\s*([\s\S]*?)
                      \s*A[\)\.:]\s*([\s\S]*?)
                      \s*B[\)\.:]\s*([\s\S]*?)
                      \s*C[\)\.:]\s*([\s\S]*?)
                      \s*D[\)\.:]\s*([\s\S]*?)
                      \s*Answer[:\s]*([ABCD])
                    """
        matches = re.findall(pattern, output_text, re.IGNORECASE | re.VERBOSE)

        quiz_data = []
        for q in matches:
            question_text, A, B, C, D, ans = q
            clean = lambda s: re.sub(r'\s+', ' ', s.strip())
            quiz_data.append({
                "question": clean(question_text),
                "options": {"A": clean(A), "B": clean(B), "C": clean(C), "D": clean(D)},
                "answer": ans.strip().upper()
            })

        if not quiz_data:
            st.warning(
                "⚠️ The model returned text but the format was unclear. Try shortening or simplifying your input.")

        return {"raw_text": raw_text, "context_text": context_text, "num_mcqs": num_mcqs, "quiz_data": quiz_data}


    # ------------------- Display Quiz Node (Updated with stylish cards) -------------------
    def display_quiz_node(state: QuizState) -> QuizState:
        quiz_data = state.get("quiz_data", [])
        if not quiz_data:
            return {"quiz_data": []}

        st.subheader("✅ Quiz Generated!")

        # Save quiz history
        if st.session_state.user:
            save_quiz(st.session_state.user["id"], quiz_data)

        # Display the quiz using the new custom HTML/CSS
        st.markdown("<span style='display:none'>.</span>", unsafe_allow_html=True)
        for i, q in enumerate(quiz_data):
            question_text = q['question'].strip()
            # The download button is created below, so we'll only display the Q&A here
            st.markdown(f"""
            <div class='quiz-card' style='animation-delay:{i * 0.2}s'>
                <div class='quiz-question'>Q{i + 1}. {question_text}</div>
                <div class='quiz-option'>A) {q['options']['A']}</div>
                <div class='quiz-option'>B) {q['options']['B']}</div>
                <div class='quiz-option'>C) {q['options']['C']}</div>
                <div class='quiz-option'>D) {q['options']['D']}</div>
                <div style='color: #00FFE0; margin-top: 15px; font-weight: bold;'>Correct Answer: {q['answer']}</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")

        # --- Download Button ---
        # Create Word Document
        doc = DocxDocument()
        doc.add_heading("AI Generated Quiz", 0)
        doc.add_paragraph("--- ANSWER KEY ---")  # Added Answer Key section

        for i, q in enumerate(quiz_data):
            doc.add_paragraph(f"Q{i + 1}. {q['question']}")
            for k, v in q["options"].items():
                doc.add_paragraph(f"{k}) {v}")
            doc.add_paragraph(f"Correct Answer: {q['answer']}\n")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        # Use the standard Streamlit download button for consistency with the new style
        st.download_button("⬇️ Download as Word", buf, "AI_Quiz.docx")

        return {"quiz_data": quiz_data}


    # ------------------- Graph Setup -------------------
    graph.add_node("upload_node", upload_node)
    graph.add_node("extract_text_node", extract_text_node)
    graph.add_node("generate_quiz_node", generate_quiz_node)
    graph.add_node("display_quiz_node", display_quiz_node)

    graph.set_entry_point("upload_node")
    graph.add_edge("upload_node", "extract_text_node")
    graph.add_edge("extract_text_node", "generate_quiz_node")
    graph.add_edge("generate_quiz_node", "display_quiz_node")
    graph.add_edge("display_quiz_node", END)

    # Compile the graph
    app = graph.compile()

    # Run App logic
    state = st.session_state.state

    # This runs the upload node to get user input for file/topic and num_mcqs
    updated = upload_node(state)
    st.session_state.state = updated  # Update state with latest user input

    if st.button("🚀 Generate Quiz"):
        # Retrieve the latest input from the state
        current_file = updated.get("file")
        current_topic = updated.get("manual_topic", "").strip()

        if not st.session_state.use_text and not current_file:
            st.error("Please upload a file or check the box to enter a topic manually.")
        elif st.session_state.use_text and not current_topic:
            st.error("Please enter a topic or text when using the manual input option.")
        else:
            with st.spinner("Generating quiz... please wait"):
                # Run the graph manually step-by-step for better Streamlit control
                updated = extract_text_node(updated)
                st.session_state.state = updated

                if updated.get("raw_text"):
                    updated = generate_quiz_node(updated)
                    st.session_state.state = updated

                    if updated.get("quiz_data"):
                        display_quiz_node(updated)
                        st.session_state.state = updated


else:
    st.markdown("""
    <style>

    /* ============================= */
    /* FIX SIDEBAR WIDTH & LOCK IT */
    /* ============================= */

    section[data-testid="stSidebar"] {
        width: 380px !important;
        min-width: 380px !important;
        max-width: 380px !important;
        resize: none !important;
    }

    section[data-testid="stSidebar"] > div:first-child {
        resize: none !important;
        overflow: auto !important;
    }

    @media (max-width: 768px) {
        section[data-testid="stSidebar"] {
            width: 300px !important;
            min-width: 300px !important;
            max-width: 300px !important;
        }
    }

    /* ============================= */
    /* SHIFTED FROM CENTER (RIGHT) */
    /* ============================= */

    .right-login-box {
        display: flex;
        justify-content: center;
        align-items: center;
        position: fixed; 
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background-color: rgba(22, 34, 49, 0.0); 
        z-index: 9999; 
        pointer-events: none;

        /* 🔥 MORE SHIFT FROM CENTER 🔥 */
        transform: translateX(200px);  
    }

    /* ============================= */
    /* ANIMATED MESSAGE */
    /* ============================= */

    .animated-message {
        background: linear-gradient(135deg, #162231, #0B1A2B);
        color: #33CCFF;
        border: 2px solid #00FFE0;

        padding: 70px 100px;
        border-radius: 22px;
        font-size: 3.0em;
        font-weight: 800;
        max-width: 90%;

        box-shadow: 
        0 0 40px rgba(0, 255, 224, 0.5),
        0 0 20px rgba(51, 204, 255, 0.8);
        text-align: center;

        animation: fadein 1s, pulse 2s infinite alternate;
    }

    /* Mobile responsive */
    @media (max-width: 800px) {
        .right-login-box {
            transform: translateX(0px);  /* Back to center on mobile */
        }

        .animated-message {
            padding: 40px 40px; 
            font-size: 2.0em;
        }
    }

    /* Animations */

    @keyframes fadein {
        from { opacity: 0; transform: scale(0.8); }
        to { opacity: 1; transform: scale(1); }
    }

    @keyframes pulse {
        0% { 
            box-shadow: 0 0 15px rgba(0, 255, 224, 0.4), 
                        0 0 5px rgba(51, 204, 255, 0.6); 
        }
        100% { 
            box-shadow: 0 0 30px rgba(0, 255, 224, 0.6), 
                        0 0 10px rgba(51, 204, 255, 0.9); 
        }
    }

    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="right-login-box">
        <div class="animated-message">
            Log-in or Sign-up
        </div>
    </div>
    """, unsafe_allow_html=True)